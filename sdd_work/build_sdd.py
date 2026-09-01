from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"D:\BCS\THIRD YEAR\PROJECT-Final Year Project\Project\WMS-FumbaPort")
TEMPLATE = ROOT / "sdd_work" / "SDD_Template_Converted.docx"
OUT = ROOT / "Fumba_Port_WMS_Software_Design_Document.docx"
WORK = ROOT / "sdd_work"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT = "D9EAF7"
PALE = "EDF4FA"
GRAY = "5B6573"
WHITE = "FFFFFF"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = tc.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tc.append(tcMar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders"); tblPr.append(borders)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        node=OxmlElement(f"w:{edge}"); node.set(qn("w:val"),"single"); node.set(qn("w:sz"),"4"); node.set(qn("w:color"),"B7C9D9"); borders.append(node)
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblGrid = table._tbl.tblGrid
    for child in list(tblGrid): tblGrid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(width)); tblGrid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            margins(cell)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true"); trPr.append(tblHeader)


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = str(h); shade(c, NAVY)
        for r in c.paragraphs[0].runs:
            r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(font_size)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    repeat_header(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            if ri % 2: shade(cells[i], PALE)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs: r.font.size = Pt(font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths is None:
        widths = [9360 // len(headers)] * len(headers); widths[-1] += 9360 - sum(widths)
    set_table_width(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="Bullet" if "Bullet" in [s.name for s in doc.styles] else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.add_run(text)
    return p


def add_number(doc, text):
    # The retained SUZA template provides a real numbering-backed Bullet style
    # but no List Number style; reuse its genuine list definition.
    p = doc.add_paragraph(style="Bullet")
    p.add_run(text)
    return p


def add_note(doc, label, text):
    t = doc.add_table(rows=1, cols=1); set_table_width(t, [9360])
    shade(t.cell(0,0), LIGHT)
    p=t.cell(0,0).paragraphs[0]; r=p.add_run(label + ": "); r.bold=True; r.font.color.rgb=RGBColor(23,54,93)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_caption(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(91,101,115)


def diagram(path, title, boxes, arrows):
    img=Image.new("RGB", (1500, 850), "white"); d=ImageDraw.Draw(img)
    try:
        font=ImageFont.truetype("arial.ttf", 28); bold=ImageFont.truetype("arialbd.ttf", 32); small=ImageFont.truetype("arial.ttf", 22)
    except Exception:
        font=bold=small=ImageFont.load_default()
    d.text((750,35), title, font=bold, fill=(23,54,93), anchor="ma")
    coords={}
    for key, label, xy in boxes:
        x1,y1,x2,y2=xy; coords[key]=xy
        d.rounded_rectangle(xy, radius=18, fill=(237,244,250), outline=(47,117,181), width=4)
        lines=label.split("\n")
        h=sum(d.textbbox((0,0), line, font=font)[3] for line in lines)+8*(len(lines)-1)
        y=(y1+y2-h)/2
        for line in lines:
            d.text(((x1+x2)/2,y), line, font=font, fill=(25,35,48), anchor="ma"); y += d.textbbox((0,0),line,font=font)[3]+8
    for a,b,label in arrows:
        A=coords[a]; B=coords[b]
        ax=(A[0]+A[2])//2; ay=(A[1]+A[3])//2; bx=(B[0]+B[2])//2; by=(B[1]+B[3])//2
        if abs(bx-ax)>abs(by-ay):
            start=(A[2] if bx>ax else A[0], ay); end=(B[0] if bx>ax else B[2], by)
        else:
            start=(ax, A[3] if by>ay else A[1]); end=(bx, B[1] if by>ay else B[3])
        d.line([start,end], fill=(91,101,115), width=4)
        import math
        ang=math.atan2(end[1]-start[1],end[0]-start[0]); L=18
        pts=[end,(end[0]-L*math.cos(ang-.5),end[1]-L*math.sin(ang-.5)),(end[0]-L*math.cos(ang+.5),end[1]-L*math.sin(ang+.5))]
        d.polygon(pts, fill=(91,101,115))
        if label: d.text(((start[0]+end[0])//2,(start[1]+end[1])//2-10),label,font=small,fill=(91,101,115),anchor="ms")
    img.save(path)


def page_break(doc): doc.add_page_break()


def clear_body(doc):
    body=doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"): body.remove(child)


def field(paragraph, instruction):
    run=paragraph.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=instruction
    sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate")
    txt=OxmlElement("w:t"); txt.text="1"
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    run._r.extend([begin,instr,sep,txt,end])


doc=Document(TEMPLATE)
clear_body(doc)
sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)

# Consistent template-derived style refinement
for style_name in ("Normal","Body Text"):
    if style_name in [s.name for s in doc.styles]:
        st=doc.styles[style_name]; st.font.name="Arial"; st.font.size=Pt(10.5); st.font.color.rgb=RGBColor(30,37,45)
        st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.08
for name,size,color in (("Title",24,NAVY),("Heading 1",16,NAVY),("Heading 2",13,BLUE),("Heading 3",11,GRAY)):
    st=doc.styles[name]; st.font.name="Arial"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(5); st.paragraph_format.keep_with_next=True
    # Template headings carry automatic outline numbers. Content headings already
    # contain the intended SDD section labels, so remove the inherited numPr to
    # prevent duplicate labels such as "4 1. INTRODUCTION".
    pPr=st._element.get_or_add_pPr(); numPr=pPr.find(qn("w:numPr"))
    if numPr is not None: pPr.remove(numPr)
if "Caption" not in [s.name for s in doc.styles]: doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

# Header/footer
hp=sec.header.paragraphs[0]; hp.text="FUMBA PORT WAREHOUSE MANAGEMENT SYSTEM  |  SOFTWARE DESIGN DOCUMENT"; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.add_run("Software Design Document  |  Page "); field(fp,"PAGE"); fp.add_run(" of "); field(fp,"NUMPAGES")
for r in fp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)

# Cover
for text,size,bold,color in [
    ("THE STATE UNIVERSITY OF ZANZIBAR",17,True,NAVY),
    ("SCHOOL OF COMPUTING, COMMUNICATION AND MEDIA STUDIES",12,True,NAVY),
    ("DEPARTMENT OF COMPUTER SCIENCE AND INFORMATION TECHNOLOGY",11,True,NAVY),
]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.font.name="Bookman Old Style"; r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
doc.add_paragraph("\n")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
r=p.add_run("SOFTWARE DESIGN DOCUMENT"); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Designing and Developing a Warehouse Management System\nfor Fumba Port Indoor Storage Facilities"); r.bold=True; r.font.size=Pt(17); r.font.color.rgb=RGBColor.from_string(BLUE)
doc.add_paragraph("\n")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Final Year Project • Bachelor of Science in Computer Science\n").bold=True
p.add_run("Academic Year 2025/2026\nVersion 1.0 • 1 September 2026")
doc.add_paragraph("\n\n")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Prepared for academic assessment and system implementation reference").italic=True
page_break(doc)

doc.add_heading("Document Control",1)
add_table(doc,["Item","Value"],[
    ("Document title","Software Design Document – Fumba Port Warehouse Management System"),
    ("System short name","Fumba Port WMS"),("Version","1.0"),("Status","Implementation-aligned baseline"),
    ("Prepared for","The State University of Zanzibar – Final Year Project"),("Date","1 September 2026"),
    ("Configuration baseline","Repository implementation and database migrations available on the document date"),
],[2200,7160])
doc.add_heading("Approval Record",2)
add_table(doc,["Role","Responsibility","Approval"],[
    ("Project author","Document preparation and implementation traceability","To be signed"),
    ("Project supervisor","Academic and technical review","To be signed"),
    ("User representative","Operational suitability review","To be signed"),
],[2200,5000,2160])
doc.add_heading("Revision History",2)
add_table(doc,["Version","Date","Description"],[("1.0","1 Sep 2026","Initial complete SDD aligned with the implemented Fumba Port WMS.")],[1200,1700,6460])
doc.add_heading("Overview",1)
doc.add_paragraph("This document defines the architecture, data design, user interfaces, module logic, integrations and integrity controls of the Fumba Port Warehouse Management System. It converts the supplied SDD template into an implementation-specific design baseline for development, testing, deployment and academic assessment.")
add_note(doc,"Scope statement","The document describes the repository implementation as of 1 September 2026. Cloud hosting, production SMTP operation and live-payment credentials are outside the demonstrated deployment baseline.")
page_break(doc)

doc.add_heading("Table of Contents",1)
for item in ["1. Introduction","2. System Architecture","3. File and Database Design","4. Human-Machine Interface","5. Detailed Design","6. External Interfaces","7. System Integrity Controls","Appendix A. Data Dictionary","Appendix B. Requirements Traceability","Appendix C. Deployment and Operations"]:
    add_bullet(doc,item)
doc.add_paragraph("Word users can update page numbering from the References tab if the document is expanded.").italic=True

doc.add_heading("1. INTRODUCTION",1)
doc.add_heading("1.1 Purpose and Scope",2)
doc.add_paragraph("The purpose of this Software Design Document (SDD) is to define how the Fumba Port WMS fulfils the operational need for secure, traceable and capacity-aware indoor freight handling. The design covers the browser user interface, REST and real-time backend, PostgreSQL persistence, Docker deployment, role-based authorization and Flutterwave payment integration.")
doc.add_paragraph("In scope are cargo registration and review; configurable warehouse hierarchy; rule-based bin recommendation; scanner-assisted placement; customs inspection; tariff, invoice and payment processing; management release; gate-out; notifications, reporting, configuration backup and immutable audit trails. Out of scope are vessel scheduling, outdoor yard planning, physical PLC/conveyor control, customs-government data exchange and public-cloud production operations.")
doc.add_heading("1.2 Project Executive Summary",2)
doc.add_paragraph("Manual freight tracking at an indoor port warehouse creates risks of lost records, over-capacity storage, delayed clearance, billing errors and unauthorized release. The WMS provides one controlled digital workflow from receipt to dispatch. Cargo is registered against an effective tariff, reviewed by a supervisor, placed only in eligible capacity, cleared by customs, financially settled or management-released, and dispatched only after a final readiness decision.")
doc.add_heading("1.2.1 System Overview",3)
doc.add_paragraph("Users access role-specific React portals through a browser or handheld scanner. Nginx terminates client connections in production and forwards API and Socket.IO traffic to an Express application. Business services enforce workflow and financial rules and persist transactional state in PostgreSQL. Flutterwave is the only external transactional service in the current design; SMTP is optional for payment notifications.")

arch=WORK/"architecture.png"
diagram(arch,"Fumba Port WMS – Logical Architecture",[
    ("users","Role portals / handheld scanner",(80,300,390,480)),
    ("edge","Nginx reverse proxy\nHTTP / HTTPS",(510,120,840,285)),
    ("api","Express API + Socket.IO\nNode.js service",(510,460,840,650)),
    ("db","PostgreSQL 17\ntransaction and audit data",(1030,470,1390,650)),
    ("pay","Flutterwave v4 / SMTP\nexternal services",(1030,120,1390,285)),
],[
    ("users","edge","HTTPS"),("edge","api","REST / WSS"),("api","db","SQL / locks"),("api","pay","OAuth / webhooks")
])
doc.add_picture(str(arch),width=Inches(6.45)); add_caption(doc,"Figure 1. High-level logical architecture")

doc.add_heading("1.2.2 Design Constraints",3)
for x in [
    "The demonstration environment is a single-host Docker Compose deployment bound to localhost by default.",
    "PostgreSQL is the authoritative data store; workflow state transitions must be transactional and concurrency-safe.",
    "Financial values must use database numeric types and fixed-point application arithmetic, never binary floating point.",
    "The system must support nine operational identities with least-privilege permissions and warehouse/shift scoping.",
    "Cargo release must not bypass customs clearance or physical placement; management release may override payment only.",
    "Uploaded documents are limited to allowed formats and validated using file signatures, not extensions alone.",
]: add_bullet(doc,x)
doc.add_heading("1.2.3 Assumptions and Future Contingencies",3)
doc.add_paragraph("The design assumes reliable local connectivity, barcode-capable scanners, trained role holders and correctly configured tariffs, warehouse hierarchy and permissions. Future growth may require managed PostgreSQL, object storage for cargo documents, a Redis-backed Socket.IO adapter, centralized observability, multi-site tenancy and a formal customs interface. The modular service and route boundaries allow those changes without replacing the user interface or domain workflow.")
doc.add_heading("1.3 Document Organization",2)
doc.add_paragraph("Section 2 defines architecture; Section 3 defines persistent and file data; Section 4 defines operator interaction; Section 5 specifies module behavior; Section 6 describes external interfaces; Section 7 defines security and integrity controls. Appendices provide a compact data dictionary, traceability matrix and deployment baseline.")
doc.add_heading("1.4 Points of Contact",2)
add_table(doc,["Contact role","Area"],[
    ("Project author / developer","Design, implementation, test evidence and configuration"),
    ("Academic supervisor","Research method, academic quality and acceptance"),
    ("Warehouse operations representative","Cargo handling, placement and dispatch validation"),
    ("Finance and customs representatives","Tariff/payment and inspection/clearance validation"),
],[3000,6360])
doc.add_heading("1.5 Project References",2)
for x in ["Repository README and source code baseline.","backend/database/schema.sql and dated SQL migrations.","Backend route, controller, service and automated-test suites.","Docker Compose, Nginx and HTTPS deployment configuration.","Flutterwave v4 OAuth/payment and webhook integration implemented by the project."]:
    add_bullet(doc,x)
doc.add_heading("1.6 Glossary",2)
add_table(doc,["Term","Meaning"],[
    ("API","Application Programming Interface"),("CRUD","Create, Read, Update and Delete"),("DBMS","Database Management System"),
    ("FRD/SRS","Functional/Software Requirements Document"),("HMAC","Hash-based Message Authentication Code"),("JWT","JSON Web Token"),
    ("RBAC","Role-Based Access Control"),("RTM","Requirements Traceability Matrix"),("SDD","Software Design Document"),
    ("WMS","Warehouse Management System"),("WSS","Secure WebSocket communication"),
],[1800,7560])

page_break(doc)
doc.add_heading("2. SYSTEM ARCHITECTURE",1)
doc.add_heading("2.1 Architectural Style and Principles",2)
doc.add_paragraph("The system uses a three-tier web architecture with a componentized single-page application, a stateless HTTP API plus stateful real-time channel, and a relational database. Controllers handle protocol concerns; services own domain rules; configuration registries provide policy authority; SQL constraints and transactions provide last-line integrity.")
for x in ["Single source of truth for workflow and authorization policy.","Defence in depth across browser, API middleware, service logic and database constraints.","Explicit public references for user-facing identities while internal keys remain private.","Idempotent handling of payment webhooks and replay-resistant refresh-token rotation.","Configuration changes are validated, audited and recoverable through backup/restore."]:
    add_bullet(doc,x)
doc.add_heading("2.2 System Hardware Architecture",2)
add_table(doc,["Component","Minimum / recommended role","Connectivity"],[
    ("Application host","64-bit dual-core CPU, 8 GB RAM, 20 GB free storage; Docker-capable OS","LAN; optional Internet for payment/email"),
    ("Operator workstation","Modern browser, 1366×768 or higher","HTTPS to application host"),
    ("Handheld scanner","Barcode reader with browser and Wi-Fi; camera or keyboard-wedge input","HTTPS/WSS over trusted WLAN"),
    ("Barcode/office printer","Standard label or A4 printer","Operator workstation / LAN"),
    ("Backup storage","Encrypted external or network storage sized to database and document retention","Scheduled administrative access"),
],[1900,4460,3000])
doc.add_heading("2.3 System Software Architecture",2)
add_table(doc,["Layer / component","Technology","Responsibility"],[
    ("Presentation","React 18, Vite, React Router, React Query, Radix/Tailwind","Role portals, forms, dashboards, responsive interaction and server-state caching"),
    ("API edge","Nginx 1.27","Static delivery, reverse proxy, security headers and TLS termination"),
    ("Application","Node.js 18+, Express 4","REST endpoints, validation, authentication and response handling"),
    ("Domain services","JavaScript service modules","Cargo workflow, placement, finance, readiness, notifications and configuration"),
    ("Real time","Socket.IO 4","Scanner sessions and operational event broadcasts"),
    ("Persistence","PostgreSQL 17 + pgcrypto","Transactions, constraints, indexed records, locks and immutable audit data"),
    ("Infrastructure","Docker / Compose","Repeatable development and production-like deployment"),
],[1750,2350,5260])
doc.add_heading("2.3.1 Backend Module Decomposition",3)
add_table(doc,["Module","Key responsibilities"],[
    ("Authentication and sessions","Login, refresh-family rotation, logout, password changes and scanner-account isolation."),
    ("Cargo and approvals","Registration, duplicate checks, documents, supervisor decision and correction/resubmission."),
    ("Warehouse hierarchy","Warehouses, zones, racks, levels, bins, shifts, capacity and assignment history."),
    ("Placement and scanning","Rule evaluation, recommendation, validation, scan pairing, placement confirmation and overrides."),
    ("Customs","Inspection start, holds, status transitions, clearance and history."),
    ("Finance and payments","Tariffs, approval, draft/issued invoices, payment recording, Flutterwave initiation and webhook settlement."),
    ("Release and gate","Readiness evaluation, management/emergency decisions and final dispatch."),
    ("Governance","RBAC, audit, notifications, reports, readiness, configuration backup/restore."),
],[2600,6760])
doc.add_heading("2.4 Internal Communications Architecture",2)
add_table(doc,["Path","Protocol / format","Design behavior"],[
    ("Browser ↔ Nginx","HTTPS; JSON; static assets","Same-origin production access; cacheable frontend assets."),
    ("Nginx ↔ backend","HTTP/1.1; WebSocket upgrade","Routes `/api` and Socket.IO traffic to the application service."),
    ("Backend ↔ PostgreSQL","TCP; PostgreSQL wire protocol","Parameterized SQL, pooled connections, transactions and advisory/row locks."),
    ("Browser ↔ Socket.IO","WSS events","Authenticated role/user rooms; scanner session feedback and operational updates."),
    ("Container network","Docker bridge DNS","Services use stable Compose service names; database is not public in production."),
],[2200,2600,4560])
add_note(doc,"Communication rule","REST is authoritative for commands and queries. Socket.IO carries notifications and live status only; clients reconcile with REST after reconnecting.")

page_break(doc)
doc.add_heading("3. FILE AND DATABASE DESIGN",1)
doc.add_heading("3.1 Database Management System Files",2)
doc.add_paragraph("PostgreSQL stores all authoritative configuration, identity, cargo, financial, inspection, dispatch and audit data. The schema is created from schema.sql and evolved by ordered migrations recorded in schema_migrations. Foreign keys preserve relationships; check constraints restrict states; indexes support queues and reporting; transactions protect multi-record workflow changes.")
doc.add_heading("3.1.1 Logical Data Model",3)
add_table(doc,["Domain","Principal tables","Relationship summary"],[
    ("Identity and access","roles, permissions, role_permissions, users, user_sessions, scanner_accounts, scanner_sessions","Users belong to a role and optional warehouse/shift; sessions and scanner pairing are independently revocable."),
    ("Storage topology","warehouses, zones, racks, levels, bins, capacity configurations","Strict warehouse→zone→rack→level→bin hierarchy with weight/volume capacity and status."),
    ("Cargo workflow","cargo, cargo_documents, approval_requests, cargo_approval_history, cargo_movements, cargo_locations","Cargo owns documents and histories; current location is separately tracked and audited."),
    ("Placement","bin_rules, placement_validation_logs, barcode_print_logs, bin_barcode_print_logs","Configurable rules determine eligible bins; every validation and print operation is recorded."),
    ("Finance","storage_tariffs, tariff_approval_requests, invoices, invoice_items, payments, payment_attempts, payment_webhook_events","Cargo is billed by an approved tariff; invoice and payment lifecycles preserve ledger history and idempotency."),
    ("Release and governance","management_release_requests, dispatch_requests, notifications, audit_logs, archived_audit_logs, system_settings","Release decisions, operator alerts, configuration and immutable evidence are retained."),
],[1700,3600,4060])
doc.add_heading("3.1.2 Key Data Rules and Access Methods",3)
for x in [
    "Cargo identity is protected by partial unique indexes covering active delivery-note, container and vehicle/consignee/type combinations.",
    "Queue access uses composite indexes on registration, placement, approval, dispatch and notification status with descending creation dates.",
    "Foreign keys bind cargo to warehouses, users, locations and workflow records; service transactions lock mutable rows before decisions.",
    "Audit tables deny UPDATE, DELETE and TRUNCATE privileges to the runtime user; archives preserve historical records.",
    "The expected data volume is moderate for a single facility. Indexes support online transaction processing; archival and backup policies manage long-term growth.",
]: add_bullet(doc,x)
doc.add_heading("3.1.3 Transaction Boundaries",3)
add_table(doc,["Transaction","Atomic effects"],[
    ("Cargo registration","Validate configuration and tariff; insert cargo; create draft invoice; create audit/notification records."),
    ("Supervisor approval","Lock cargo/invoice; record decision; issue invoice and public reference; queue customer notification."),
    ("Placement confirmation","Lock cargo/bin; re-evaluate rules and capacity; write location/movement/log; update occupied capacity."),
    ("Payment settlement","Verify invoice and idempotency key; write payment/attempt event; update totals and status."),
    ("Gate-out","Re-evaluate release readiness; lock cargo; record dispatch; set placement state to Dispatched; audit actor."),
],[2500,6860])
doc.add_heading("3.2 Non-DBMS Files",2)
add_table(doc,["File class","Purpose","Readers / writers","Control"],[
    ("Cargo documents","Operator-uploaded supporting evidence","Cargo module; authorized reviewers","Magic-byte/type/size validation; generated safe name; access-controlled retrieval"),
    ("Configuration backup","Portable administrative snapshot","System administrator","Schema and content validation before restore; operation audit and rate limit"),
    ("Report export","CSV/PDF-like downloadable operational output","Authorized finance, management and role-report users","Generated on demand; no authoritative state"),
    ("Application logs","Operational diagnosis and email fallback","Backend runtime / administrator","Sensitive-value redaction; environment-controlled retention"),
    ("Frontend assets","Compiled HTML, CSS, JavaScript and images","Nginx / browsers","Immutable build output and cache policy"),
],[1700,2300,2300,3060])

page_break(doc)
doc.add_heading("4. HUMAN-MACHINE INTERFACE",1)
doc.add_heading("4.1 Interaction Model",2)
doc.add_paragraph("The interface presents a landing/login flow followed by role-specific portals. A shared application layout provides navigation, authenticated identity, notifications and logout. Permissions from the backend hide or disable actions, but server authorization remains decisive. Destructive or workflow-changing operations use confirmation dialogs and require notes where accountability is needed.")
doc.add_heading("4.2 Inputs",2)
add_table(doc,["Input / screen","Primary users","Core data and validation"],[
    ("Login and initial setup","All roles / initial administrator","Username/password; rate limit; one-time bootstrap lock; password policy."),
    ("Cargo registration","Warehouse staff","Consignee, delivery/container/vehicle identifiers, cargo type, dimensions, weight, hazard/fragility, warehouse, documents; mandatory and duplicate validation."),
    ("Supervisor review","Supervisor","Approve, reject or request correction; decision notes required for adverse/exception outcomes."),
    ("Scanner placement","Staff and paired scanner","Cargo barcode then bin barcode; active session, assignment, capacity, rule and status validation."),
    ("Customs processing","Customs officer","Inspection start, clearance/hold status and notes; authorized transitions only."),
    ("Tariff and payment","Finance / customer","Rate basis, effective dates, approval; payment initiation through time-bound public token; idempotent confirmation."),
    ("Release / gate-out","Management / gate","Release request reasons, decision notes, cargo barcode/reference; readiness rechecked at confirmation."),
    ("System configuration","Administrator","Hierarchy, capacity, roles, permissions, forms, notifications and backup; validation and audit required."),
],[2150,1850,5360])
doc.add_heading("4.2.1 Input Control Rules",3)
for x in [
    "Client-side Zod/form checks provide immediate guidance; API validation rejects bypassed or malformed requests.",
    "Identifiers, numeric bounds, state enums, object depth and property counts are constrained.",
    "Document content is validated by binary signature and allowed MIME/type policy.",
    "Permission and warehouse scope are derived from authenticated context, not accepted from untrusted client claims.",
    "Barcode operations bind to a live scanner session and server-side cargo/bin records.",
]: add_bullet(doc,x)
doc.add_heading("4.3 Outputs",2)
add_table(doc,["Output","Audience","Purpose and access"],[
    ("Role dashboards","Each operational role","Queues, counts, alerts and next actions limited by permission and warehouse scope."),
    ("Barcode labels","Warehouse staff","Machine-readable cargo/bin identifiers with human-readable reference; print event logged."),
    ("Invoice and payment link","Finance/customer","Itemized charges, paid amount, balance and approved online-payment action."),
    ("Placement recommendation","Warehouse staff/supervisor","Ranked eligible bins and rule explanations; recommendation is revalidated at placement."),
    ("Customs/release history","Customs, gate, supervisor, auditor","Chronological evidence supporting clearance and release decisions."),
    ("Operational and financial reports","Management, finance, auditor","Filtered summaries and exports controlled by report permissions."),
    ("Notifications","Authorized recipients","Pending reviews, escalations, announcements and workflow events via portal/realtime channel."),
    ("Audit log","Administrator/auditor","Actor, time, network context, action, target and change metadata."),
],[2400,2100,4860])
doc.add_heading("4.4 Accessibility and Usability",2)
doc.add_paragraph("Interfaces use semantic labels, keyboard-capable component primitives, visible focus states, descriptive feedback and responsive layouts. Status is communicated through text as well as color. Tables provide headings and filters; dialogs keep action labels explicit. The production acceptance test should include keyboard-only use, contrast review, screen-reader labels and handheld viewport testing.")

page_break(doc)
doc.add_heading("5. DETAILED DESIGN",1)
doc.add_heading("5.1 Hardware Detailed Design",2)
doc.add_paragraph("The application does not directly control industrial hardware. It uses standard computers, printers and browser-capable scanners. Production procurement should select devices that satisfy the following logical requirements.")
add_table(doc,["Item","Detailed requirement"],[
    ("Server","64-bit CPU, at least 8 GB RAM, SSD storage, reliable clock synchronization, daily backup target and UPS protection."),
    ("Client","Current Chromium/Firefox/Edge browser, 1366×768 display or better, keyboard and pointing device."),
    ("Scanner","Reads configured 1D/2D barcode format; injects text or exposes browser camera; stable Wi-Fi; protected user session."),
    ("Printer","Supports the selected cargo/bin label size and A4 invoice/report output; client OS driver supplied by vendor."),
    ("Network","Trusted LAN/WLAN, TLS-capable edge, firewall allowing only required application ports; Internet egress for Flutterwave/SMTP."),
],[1800,7560])
doc.add_heading("5.2 Software Detailed Design",2)
doc.add_heading("5.2.1 Authentication and Authorization",3)
doc.add_paragraph("The login controller validates credentials with bcrypt and issues a short-lived JWT plus database-backed refresh session. Refresh tokens rotate by family; reuse of a superseded token indicates replay and revokes the affected family. Middleware authenticates the request, distinguishes scanner accounts and checks a centralized permission registry. Services additionally apply warehouse/role scope to data queries.")
doc.add_heading("5.2.2 Cargo Registration and Review",3)
add_number(doc,"Load the published registration-form configuration and effective, management-approved tariff.")
add_number(doc,"Validate cargo fields, dimensions, weights, identifiers, warehouse assignment and uploaded evidence.")
add_number(doc,"Check active duplicate identities and insert cargo in Pending Review state.")
add_number(doc,"Create a Draft invoice without a payment reference and write the audit trail.")
add_number(doc,"Supervisor approves, rejects or requests correction. Approval issues the existing invoice; rejection cancels payment capability while preserving history.")
doc.add_heading("5.2.3 Placement Rule Engine",3)
doc.add_paragraph("The bin-rule engine obtains enabled rule definitions from configuration, resolves each evaluator from an allow-listed registry and evaluates candidate bins against cargo and topology context. Typical evaluators include remaining weight/volume, cargo type, hazardous classification, fragility, customs conditions and status. Recommendations contain eligible locations and reasons; confirmation repeats evaluation inside a transaction to prevent stale-capacity placement.")
add_note(doc,"Concurrency","Placement locks the affected cargo and bin/capacity rows before writing the current location and occupied capacity. A recommendation alone never reserves capacity.")
doc.add_heading("5.2.4 Cargo State Machine",3)
workflow=WORK/"workflow.png"
diagram(workflow,"Cargo Lifecycle and Release Gates",[
    ("reg","Registered\nPending Review",(50,110,335,250)),("app","Supervisor\nApproved",(440,110,725,250)),
    ("place","Placed in\neligible bin",(830,110,1115,250)),("clear","Customs\nCleared",(1170,110,1450,250)),
    ("pay","Paid or approved\nmanagement release",(740,500,1070,650)),("gate","Gate-out\nDispatched",(1160,500,1450,650)),
    ("reject","Rejected /\nCorrection Required",(260,500,590,650)),
],[
    ("reg","app","review"),("app","place","scan + validate"),("place","clear","inspection"),("clear","pay","readiness"),("pay","gate","final check"),("reg","reject","decision")
])
doc.add_picture(str(workflow),width=Inches(6.45)); add_caption(doc,"Figure 2. Cargo workflow and mandatory release gates")
doc.add_heading("5.2.5 Customs, Finance and Release",3)
doc.add_paragraph("Customs operations record inspection history and a current clearance/hold status. Finance calculates charges with scaled BigInt arithmetic, manages tariff approval and invoice lifecycles, and accepts manual or Flutterwave settlement. Release readiness is a pure policy evaluation over registration approval, physical placement, customs clearance, invoice settlement or approved management release, and absence of blocking state. Gate-out repeats the evaluation transactionally and records the dispatch actor and time.")
doc.add_heading("5.2.6 Notifications, Reports and Audit",3)
doc.add_paragraph("Domain events create recipient-scoped notifications according to configurable policies and escalation settings. Role reports aggregate only authorized scopes and export through dedicated endpoints. Audit middleware and services capture user, role/warehouse snapshots, IP or terminal context, action, target, timestamp and metadata. Runtime database grants make audit records append-only.")
doc.add_heading("5.2.7 Error Handling and Logging",3)
doc.add_paragraph("Controllers pass failures to a common error middleware that converts known API errors to minimized JSON responses and prevents stack or secret leakage. Request identifiers and structured logs support diagnosis. Validation failures use 4xx responses; authorization uses 401/403; conflicts such as duplicate or stale state use 409; unexpected errors use 500 and are logged server-side.")
doc.add_heading("5.3 Internal Communications Detailed Design",2)
add_table(doc,["Concern","Detailed design"],[
    ("API format","JSON request/response under `/api`; success payloads contain data/count where applicable; errors are minimized."),
    ("Authentication","Bearer access token; refresh endpoint uses a rotatable server-tracked session token; scanner credentials are isolated."),
    ("Real-time events","Socket.IO authenticated connection joins permitted user/role/session rooms; events signal change and clients refresh authoritative data."),
    ("Database access","`pg` connection pool; parameterized statements; explicit BEGIN/COMMIT/ROLLBACK for multi-write operations."),
    ("Timeout/retry","External payment calls use bounded failure handling; webhook retries are safe because event/provider references are idempotent."),
],[2200,7160])

page_break(doc)
doc.add_heading("6. EXTERNAL INTERFACES",1)
doc.add_heading("6.1 Interface Architecture",2)
doc.add_paragraph("External interfaces are isolated behind backend services. Browsers never receive provider credentials or connect directly to the database. Outbound payment and email calls originate from the backend. Inbound payment callbacks terminate at a dedicated webhook route and are authenticated before any financial state changes.")
doc.add_heading("6.2 Interface Detailed Design",2)
add_table(doc,["Interface","Direction and protocol","Data / handshake","Failure handling"],[
    ("Flutterwave OAuth/payment","Backend ↔ provider over HTTPS","Client credentials obtain access token; payment request carries amount, currency, reference and return context.","Reject unavailable/invalid provider response; persist attempt state; show safe user message."),
    ("Flutterwave webhook","Provider → backend HTTPS POST","Raw payload plus configured HMAC-SHA256 signature; provider event/reference identifies transaction.","Reject invalid signature; record event idempotently; duplicate delivery returns without duplicate payment."),
    ("SMTP","Backend → mail server using TLS as configured","Payment-link and notification message to validated recipient.","Delivery failure is logged/queued; when SMTP is absent the demonstration uses logger fallback."),
    ("Browser client","Browser ↔ Nginx over HTTPS/WSS","JSON API and Socket.IO frames; access token and server permission checks.","Standard 4xx/5xx response; reconnect then REST reconciliation for real-time channel."),
],[1550,2400,3160,2250],8)
doc.add_heading("6.2.1 Public Payment Interface",3)
doc.add_paragraph("The customer opens `/pay/:token`. The server resolves a high-entropy token to an eligible issued invoice and exposes only the summary necessary to pay. Initiation is rate limited and records a payment attempt. A token is not an authentication credential for finance administration and cannot access cargo documents, internal notes or unrelated invoices.")
doc.add_heading("6.2.2 Webhook Processing Sequence",3)
for x in [
    "Capture the exact request body and required provider signature header.",
    "Compute HMAC using the configured webhook secret and compare safely.",
    "Validate event type, provider reference, currency, amount and expected invoice relationship.",
    "Insert/lock the webhook event using its unique identity; if already processed, return success without reapplying.",
    "Write payment/ledger changes and invoice status in one transaction; emit audit/notification after commit.",
]: add_number(doc,x)

doc.add_heading("7. SYSTEM INTEGRITY CONTROLS",1)
doc.add_heading("7.1 Security Control Objectives",2)
add_table(doc,["Objective","Implemented design control"],[
    ("Identification and authentication","bcrypt password hashes; short-lived JWT; refresh-session rotation, revocation and replay detection; one-time bootstrap lock."),
    ("Least privilege","Central permission registry, route middleware, non-scanner restrictions and warehouse/shift data scoping."),
    ("Confidentiality","TLS at the production edge; secret environment variables; response minimization; token/password redaction."),
    ("Integrity","Schema constraints, parameterized SQL, transactions, locks, fixed-point money arithmetic and idempotent provider events."),
    ("Availability","Container health/restart strategy, database persistence, backup/restore, bounded rate limits and operational readiness checks."),
    ("Accountability","Append-only audit data containing actor, role/warehouse snapshot, time, network context, target and metadata."),
],[2200,7160])
doc.add_heading("7.2 Input, Session and API Controls",2)
for x in [
    "Request validation rejects excessive nesting/property counts and prototype-pollution keys.",
    "Uploads are constrained by type, size and binary signature; retrieval requires authorization.",
    "Authentication, refresh, bootstrap, public payment and administrative operations have purpose-specific rate limits.",
    "Security headers, restrictive CORS and production HTTPS reduce browser attack surface.",
    "Responses exclude hashes, raw tokens, provider secrets and avoid revealing internal error stacks.",
]: add_bullet(doc,x)
doc.add_heading("7.3 Workflow and Financial Integrity",2)
doc.add_paragraph("Every critical transition is authorized, state-checked and audited. Cargo cannot be placed before approval, capacity cannot become negative, hazardous cargo cannot enter an ineligible bin, customs holds block release, and an unpaid invoice requires a separately approved management release. Management release does not override customs or placement. Gate-out is the final transactional enforcement point.")
doc.add_heading("7.4 Audit, Retention and Recovery",2)
doc.add_paragraph("Audit records are immutable to the runtime account and can be archived only through an authorized, rate-limited administrative workflow. Configuration backup includes validation before restore. Production operation should add scheduled PostgreSQL backups, encrypted off-host copies, restore drills, log retention thresholds and documented recovery time/recovery point objectives.")
doc.add_heading("7.5 Verification Strategy",2)
add_table(doc,["Verification level","Evidence"],[
    ("Unit and service","Node test suites for rule evaluators, workflow, finance, notification, configuration and security logic."),
    ("API integration","Tests for RBAC, route contracts, payment/webhook behavior, customs/gate policies and concurrency."),
    ("Frontend component","Vitest and Testing Library tests for portals, access rules, scanner/payment flows and critical components."),
    ("Database","Schema verification, ordered migrations, constraints/grant checks and PostgreSQL typing tests."),
    ("Release/UAT","Readiness service, Docker build, role-based end-to-end scenarios and localhost user acceptance testing."),
],[2200,7160])

page_break(doc)
doc.add_heading("APPENDIX A. COMPACT DATA DICTIONARY",1)
add_table(doc,["Entity","Representative elements","Validation / maintenance"],[
    ("users","public reference, username, password hash, role, warehouse, shift, status","Unique identity; administrator-managed; self-profile changes limited."),
    ("cargo","reference, consignee, delivery/container/vehicle IDs, type, weight, volume, hazard/fragility, registration/customs/placement states","Required fields and enums; duplicate checks; workflow services update."),
    ("bins","barcode, hierarchy keys, weight/volume capacity, status, special attributes","Unique within level; capacity and eligibility protected during placement."),
    ("approval_requests","cargo, type, status, requester/assignee, decision notes/times","Created by workflow; authorized decision; history retained."),
    ("invoices","number, cargo, tariff, subtotal/total, paid amount, status, public payment token","Unique number/token; fixed-point calculation; lifecycle restricted."),
    ("payments","reference, invoice, amount, method/provider, status, confirmed time","Positive amount; idempotent provider reference; finance/provider updates."),
    ("audit_logs","actor/snapshots, action, target, network context, metadata, time","Append-only runtime access; archive through controlled operation."),
    ("notifications","recipient user/role/warehouse, event type, priority, read/archive/resolve state","Policy-driven creation; user-scoped maintenance."),
],[1400,4800,3160],8)

doc.add_heading("APPENDIX B. REQUIREMENTS TRACEABILITY MATRIX",1)
add_table(doc,["ID","Requirement","Design allocation","Verification"],[
    ("FR-01","Register and review cargo","5.2.2; cargo/approval modules","Cargo workflow and supervisor tests"),
    ("FR-02","Prevent storage over-allocation","3.1.3; 5.2.3","Placement, capacity and concurrency tests"),
    ("FR-03","Barcode-assisted placement","4.2; 5.2.3","Scanner session and placement tests"),
    ("FR-04","Customs inspection and hold","5.2.5","Customs workflow/authority tests"),
    ("FR-05","Tariff, invoice and payment","3.1; 5.2.5; 6.2","Finance and Flutterwave tests"),
    ("FR-06","Controlled cargo release","5.2.4–5.2.5","Readiness, management and gate tests"),
    ("FR-07","Role-restricted portals","2.3; 7.1–7.2","RBAC and portal-access tests"),
    ("FR-08","Immutable audit and reporting","5.2.6; 7.4","Audit grants/log access/report tests"),
    ("NFR-01","Security and privacy","7.1–7.4","Security-hardening tests and review"),
    ("NFR-02","Reliability under concurrency","3.1.3; 5.2.3","Final concurrency validation"),
],[850,2650,3200,2660],8)

doc.add_heading("APPENDIX C. DEPLOYMENT AND OPERATIONS",1)
add_table(doc,["Service","Development binding","Container port","Dependency"],[
    ("Frontend","127.0.0.1:3000","3000","Backend API"),("Backend","127.0.0.1:5000","5000","PostgreSQL; optional Flutterwave/SMTP"),("PostgreSQL","127.0.0.1:5433","5432","Persistent volume"),
],[1800,2500,1900,3160])
doc.add_heading("C.1 Deployment Procedure",2)
for x in ["Create the environment file from the supplied example and replace all secrets.","Build and start the Docker Compose services.","Apply database migrations and verify schema/grants.","Run automated backend and frontend tests.","Complete one-time administrator setup and configure roles, hierarchy, tariffs and notification policies.","Execute role-based UAT for registration, placement, customs, payment and gate-out.","For production, configure DNS/TLS, backups, SMTP, provider credentials and monitoring before exposing the service."]:
    add_number(doc,x)
doc.add_heading("C.2 Known Limitations and Planned Enhancements",2)
for x in ["Current baseline targets a local single-instance demonstration, not a public multi-node deployment.","Email falls back to application logging when SMTP is not configured.","External customs integration and object storage are not implemented.","Production readiness requires penetration testing, backup restoration evidence, accessibility review and live-provider certification."]:
    add_bullet(doc,x)

# document metadata and field refresh
doc.core_properties.title="Software Design Document – Fumba Port Warehouse Management System"
doc.core_properties.subject="Final Year Project software architecture and detailed design"
doc.core_properties.author="Fumba Port WMS Project"
settings=doc.settings._element
if settings.find(qn("w:updateFields")) is None:
    u=OxmlElement("w:updateFields"); u.set(qn("w:val"),"true"); settings.append(u)

doc.save(OUT)
print(OUT)
